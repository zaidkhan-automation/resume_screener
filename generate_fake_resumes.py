"""
generate_fake_resumes.py
Generate fake resumes (PDF + DOCX) for testing your Resume Screener app.

Usage:
    python generate_fake_resumes.py --count 20

This will create out/resumes_docx/.docx and out/resumes_pdf/.pdf
"""
import os
import random
import argparse
from faker import Faker
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

fake = Faker()

SKILLS_POOL = [
    "Python","FastAPI","Django","Flask","Streamlit","Pandas","NumPy","SQL","PostgreSQL",
    "MongoDB","Docker","Kubernetes","AWS","GCP","Azure","React","Node.js","JavaScript",
    "TypeScript","Excel","Power Query","Power BI","Tableau","NLP","OCR","OpenAI","PyTorch",
    "TensorFlow","HTML","CSS","Git","Linux","Shell"
]

EDUCATION_POOL = [
    "B.Tech in Computer Science", "B.Sc in Computer Science", "M.Tech in CS", "MBA (Operations)",
    "B.Com", "MCA", "BBA", "M.Sc Data Science"
]

JOB_TITLES = [
    "Software Engineer", "Data Engineer", "ML Engineer", "Backend Developer",
    "Automation Engineer", "Data Analyst", "Business Analyst", "DevOps Engineer",
    "Technical Lead", "Fullstack Developer"
]

COMPANIES = [
    "ABC Solutions", "Delta Tech", "Innova Systems", "PixelWorks", "Nimbus Labs",
    "QuantumSoft", "HexaCorp", "BlueWave"
]

def random_skills(n=6):
    return random.sample(SKILLS_POOL, k=min(n, len(SKILLS_POOL)))

def make_person():
    name = fake.name()
    email = fake.email()
    phone = fake.phone_number()
    city = fake.city()
    degree = random.choice(EDUCATION_POOL)
    years_exp = random.randint(0, 12)
    # Build a small experience block
    exp_count = random.randint(1, 3)
    experiences = []
    for i in range(exp_count):
        title = random.choice(JOB_TITLES)
        company = random.choice(COMPANIES)
        start_year = 2024 - (years_exp - i*2 if years_exp>i*2 else random.randint(1,3))
        end_year = start_year + random.randint(1, 3)
        experiences.append({
            "title": title,
            "company": company,
            "period": f"{start_year} - {end_year}",
            "bullets": [
                f"Worked on {random.choice(SKILLS_POOL)} based projects",
                f"Improved process efficiency by {random.randint(10,40)}%",
                f"Built automation tooling for reporting"
            ]
        })
    skills = random_skills(n=random.randint(4,10))
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "city": city,
        "degree": degree,
        "years": years_exp,
        "experiences": experiences,
        "skills": skills
    }

def write_docx(person, out_path):
    doc = Document()
    doc.add_heading(person["name"], level=1)
    doc.add_paragraph(f"{person['email']} | {person['phone']} | {person['city']}")
    doc.add_paragraph(person["degree"] + f" | {person['years']} yrs experience")
    doc.add_heading("Experience", level=2)
    for e in person["experiences"]:
        doc.add_heading(f"{e['title']} — {e['company']}", level=3)
        doc.add_paragraph(e["period"])
        for b in e["bullets"]:
            doc.add_paragraph(b, style='List Bullet')
    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(person["skills"]))
    doc.save(out_path)

def write_pdf(person, out_path):
    c = canvas.Canvas(out_path, pagesize=A4)
    width, height = A4
    margin = 50
    y = height - margin
    c.setFont("Helvetica-Bold", 16)
    c.drawString(margin, y, person["name"])
    y -= 22
    c.setFont("Helvetica", 10)
    c.drawString(margin, y, f"{person['email']} | {person['phone']} | {person['city']}")
    y -= 18
    c.drawString(margin, y, f"{person['degree']} | {person['years']} yrs")
    y -= 24

    c.setFont("Helvetica-Bold", 12)
    c.drawString(margin, y, "Experience")
    y -= 16
    c.setFont("Helvetica", 10)
    for e in person["experiences"]:
        if y < 100:
            c.showPage()
            y = height - margin
        c.setFont("Helvetica-Bold", 11)
        c.drawString(margin, y, f"{e['title']} — {e['company']}")
        y -= 14
        c.setFont("Helvetica-Oblique", 9)
        c.drawString(margin, y, e["period"])
        y -= 12
        c.setFont("Helvetica", 9)
        for b in e["bullets"]:
            if y < 80:
                c.showPage()
                y = height - margin
            c.drawString(margin+10, y, f"• {b}")
            y -= 12
        y -= 8

    if y < 140:
        c.showPage()
        y = height - margin

    c.setFont("Helvetica-Bold", 12)
    c.drawString(margin, y, "Skills")
    y -= 16
    c.setFont("Helvetica", 10)
    c.drawString(margin, y, ", ".join(person["skills"]))
    c.showPage()
    c.save()

def ensure_dirs(base="out"):
    d1 = os.path.join(base, "resumes_docx")
    d2 = os.path.join(base, "resumes_pdf")
    os.makedirs(d1, exist_ok=True)
    os.makedirs(d2, exist_ok=True)
    return d1, d2

def main(count=10, out="out"):
    docx_dir, pdf_dir = ensure_dirs(out)
    for i in range(count):
        p = make_person()
        safe_name = p["name"].lower().replace(" ", "_").replace(".", "")
        docx_path = os.path.join(docx_dir, f"{safe_name}_{i+1}.docx")
        pdf_path = os.path.join(pdf_dir, f"{safe_name}_{i+1}.pdf")
        write_docx(p, docx_path)
        write_pdf(p, pdf_path)
        print(f"Generated: {docx_path}  |  {pdf_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--count", type=int, default=10, help="Number of resumes to generate")
    parser.add_argument("--out", type=str, default="out", help="Output folder")
    args = parser.parse_args()
    main(count=args.count, out=args.out)
